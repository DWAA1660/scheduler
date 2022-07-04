from docx import Document
import motor
import motor.motor_asyncio
from docx.shared import Inches
from CONFIG import CLIENT
db = CLIENT.main 
async def invoice(job_id, job_name, amount_of_shifts, shift_results, price_per_hour):
    total_price = 0.0
    total_hours = 0.0
    total_minutes = 0.0
    document = Document()
    document.add_heading(f'Invoice for {job_name}', 0)
    #time sheet
    shift_table = document.add_table(rows = 1, cols = 2)
    collum = shift_table.rows[0].cells
    collum[0].text = 'Employee Name'
    collum[1].text = 'Time worked'
    total_hours = 0
    total_minutes = 0
    async for shift in shift_results:
        employee_id = shift['employee']
        employee_results = await db.main.employee.find_one({'id': employee_id})
        row = shift_table.add_row().cells
        row[0].text = employee_results['name']
        row[1].text = shift['total_time']

        shift_minutes = shift['total_time'].split(':')[1]
        shift_hours = shift['total_time'].split(':')[0]

        total_hours = total_hours + int(shift_hours)
        total_minutes = total_minutes +  int(shift_minutes)

    total_hours = (total_minutes // 60) + total_hours
    total_minutes = total_minutes % 60

    total_time = f"{str(total_hours)}:{str(total_minutes)}"

    document.add_paragraph('Totals')

    #total price sheet

    price_per_minute = price_per_hour /60

    total_price = (total_hours * 60) * price_per_minute
    total_price = total_price +  total_minutes * price_per_minute

    total_price = round(total_price, 2)

    total_price_table = document.add_table(rows = 1, cols = 2)
    collum = total_price_table.rows[0].cells
    collum[0].text = 'Total time'
    collum[1].text = total_time 

    row = total_price_table.add_row().cells
    row[0].text = 'Price per hour'
    row[1].text = f"${float(price_per_hour)}"


    row = total_price_table.add_row().cells
    row[0].text = 'Total price'
    row[1].text = f"${float(total_price)}"


    document.save(f'static/{job_id}-demo.docx')